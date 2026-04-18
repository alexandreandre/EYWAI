"""
Moteur d'injection et conversion documents (docx / html → PDF).

Dépendances optionnelles : docxtpl, python-docx ; LibreOffice en subprocess pour docx→pdf.
"""

from __future__ import annotations

import io
import logging
import os
import re
import shutil
import subprocess
import tempfile
import zipfile
from typing import Any, Dict, List

from weasyprint import HTML

from app.services.document_variables import get_unknown_variables

logger = logging.getLogger(__name__)

_PLACEHOLDER_RE = re.compile(r"\{\{\s*([a-zA-Z0-9_]+)\s*\}\}")


class DocumentEngine:
    """Injection de variables et conversions vers PDF."""

    def inject_docx(self, template_bytes: bytes, variables: Dict[str, str]) -> bytes:
        try:
            from docxtpl import DocxTemplate
        except ImportError as e:
            raise RuntimeError(
                "docxtpl est requis pour inject_docx. Installez : pip install docxtpl"
            ) from e

        tpl = DocxTemplate(io.BytesIO(template_bytes))
        tpl.render(variables)
        out = io.BytesIO()
        tpl.save(out)
        return out.getvalue()

    def inject_html(self, template_content: str, variables: Dict[str, str]) -> str:
        def _repl(match: re.Match[str]) -> str:
            key = match.group(1).strip()
            return variables.get(key, "")

        return _PLACEHOLDER_RE.sub(_repl, template_content or "")

    def html_to_pdf(self, html_content: str) -> bytes:
        return HTML(string=html_content or "").write_pdf()

    def docx_to_pdf(self, docx_bytes: bytes) -> bytes:
        pdf_via_lo = self._docx_to_pdf_libreoffice(docx_bytes)
        if pdf_via_lo:
            return pdf_via_lo
        return self._docx_to_pdf_reportlab_fallback(docx_bytes)

    def _docx_to_pdf_libreoffice(self, docx_bytes: bytes) -> bytes | None:
        for bin_name in ("libreoffice", "soffice"):
            if not shutil.which(bin_name):
                continue
            tmp = tempfile.mkdtemp(prefix="eywai_docx_")
            try:
                docx_path = os.path.join(tmp, "template.docx")
                with open(docx_path, "wb") as f:
                    f.write(docx_bytes)
                cmd = [
                    bin_name,
                    "--headless",
                    "--norestore",
                    "--nolockcheck",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    tmp,
                    docx_path,
                ]
                proc = subprocess.run(
                    cmd,
                    capture_output=True,
                    text=True,
                    timeout=120,
                    check=False,
                )
                if proc.returncode != 0:
                    logger.warning(
                        "%s conversion failed rc=%s stderr=%s",
                        bin_name,
                        proc.returncode,
                        (proc.stderr or "")[:500],
                    )
                    continue
                pdf_path = os.path.join(tmp, "template.pdf")
                if os.path.isfile(pdf_path):
                    with open(pdf_path, "rb") as pf:
                        return pf.read()
            except (OSError, subprocess.SubprocessError) as e:
                logger.warning("LibreOffice conversion error: %s", e)
            finally:
                shutil.rmtree(tmp, ignore_errors=True)
        return None

    def _docx_to_pdf_reportlab_fallback(self, docx_bytes: bytes) -> bytes:
        try:
            from docx import Document as DocxDocument
        except ImportError as e:
            raise RuntimeError(
                "Conversion docx→pdf : installez LibreOffice, ou pip install python-docx "
                "pour le repli ReportLab."
            ) from e

        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas

        doc = DocxDocument(io.BytesIO(docx_bytes))
        lines: List[str] = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                lines.append(t)
        for table in doc.tables:
            for row in table.rows:
                cells = " | ".join(
                    (c.text or "").strip() for c in row.cells if (c.text or "").strip()
                )
                if cells:
                    lines.append(cells)

        buf = io.BytesIO()
        c = canvas.Canvas(buf, pagesize=A4)
        width, height = A4
        y = height - 40
        c.setFont("Helvetica", 9)
        for line in lines:
            if y < 40:
                c.showPage()
                c.setFont("Helvetica", 9)
                y = height - 40
            # Découpage rudimentaire
            text = line[:120]
            c.drawString(40, y, text)
            y -= 12
        c.save()
        return buf.getvalue()

    def _extract_text_from_docx_xml(self, template_bytes: bytes) -> str:
        try:
            zf = zipfile.ZipFile(io.BytesIO(template_bytes))
            with zf.open("word/document.xml") as xf:
                return xf.read().decode("utf-8", errors="ignore")
        except (KeyError, OSError, zipfile.BadZipFile):
            return ""

    def preview_variables(
        self,
        template_bytes: bytes,
        file_format: str,
        known_variables: Dict[str, str],
    ) -> Dict[str, Any]:
        fmt = (file_format or "").lower().strip()
        content = ""
        if fmt == "docx":
            content = self._extract_text_from_docx_xml(template_bytes)
        elif fmt == "html":
            content = (template_bytes or b"").decode("utf-8", errors="ignore")
        else:
            return {"unknown_variables": [], "preview_available": False}

        unknown = get_unknown_variables(content, known_variables)
        return {
            "unknown_variables": unknown,
            "preview_available": True,
        }


document_engine = DocumentEngine()
