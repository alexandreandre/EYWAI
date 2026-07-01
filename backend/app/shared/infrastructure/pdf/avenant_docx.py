"""
Génération Word (.docx) d'avenant au contrat de travail — équivalent éditable du PDF.

Même contenu que avenant.py (mentions obligatoires, tableau ancien/nouveau),
mais produit un fichier Word que le service RH peut rouvrir et retravailler.
"""

from __future__ import annotations

from datetime import date
from io import BytesIO
from typing import Any, Dict, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

from app.shared.infrastructure.pdf.avenant import (
    AVENANT_TYPE_LABELS,
    _ctx_str,
    _is_economic_motif,
)
from app.shared.infrastructure.pdf.helpers import (
    format_salary_euros,
    get_company_address,
    get_company_city,
    get_company_name,
    get_company_signatory,
    get_employee_address,
    get_lieu_travail,
    get_salary_amount,
    resolve_company_logo,
    safe_str,
)


def _fmt_date(value: Any) -> str:
    if not value:
        return "…………………"
    if isinstance(value, str):
        try:
            return date.fromisoformat(value[:10]).strftime("%d/%m/%Y")
        except ValueError:
            return value
    if isinstance(value, date):
        return value.strftime("%d/%m/%Y")
    return str(value)


def _fmt_salary(value: Any) -> str:
    if value is None or value == "":
        return "…………………"
    if isinstance(value, dict):
        return format_salary_euros({"salaire_de_base": value})
    try:
        return format_salary_euros({"salaire_de_base": {"valeur": float(value)}})
    except (TypeError, ValueError):
        return safe_str(value)


def _fmt_duree(value: Any) -> str:
    if value is None or value == "":
        return "…………………"
    try:
        f = float(str(value).replace(",", ".").replace(" h", "").strip())
        return f"{f:g} heures par semaine".replace(".", ",")
    except (ValueError, TypeError):
        s = safe_str(value)
        return s if "h" in s.lower() else f"{s} heures par semaine"


def _add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(14)
    p.paragraph_format.space_after = Pt(14)


def _add_article_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11.5)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def _add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text)
    p.paragraph_format.space_after = Pt(4)


def _add_modif_table(
    doc: Document,
    avenant_type: str,
    context: Dict[str, Any],
    employee_data: Dict[str, Any],
) -> None:
    motif = _ctx_str(context, "motif_avenant", "motif")

    if avenant_type == "avenant_salaire":
        ancien = _ctx_str(context, "ancien_salaire") or str(
            get_salary_amount(employee_data) or ""
        )
        nouveau = _ctx_str(context, "nouveau_salaire")
        _render_table(doc, "Rémunération mensuelle brute", _fmt_salary(ancien), _fmt_salary(nouveau))
        return

    if avenant_type == "avenant_poste":
        ancien = _ctx_str(context, "ancien_poste") or safe_str(
            employee_data.get("job_title", "")
        )
        nouveau = _ctx_str(context, "nouveau_poste")
        _render_table(doc, "Fonctions / emploi", ancien or "…………………", nouveau or "…………………")
        return

    if avenant_type == "avenant_temps":
        ancien = _ctx_str(context, "ancienne_duree") or safe_str(
            employee_data.get("duree_hebdomadaire", "")
        )
        nouveau = _ctx_str(context, "nouvelle_duree")
        _render_table(doc, "Durée hebdomadaire de travail", _fmt_duree(ancien), _fmt_duree(nouveau))
        return

    if avenant_type == "avenant_lieu":
        ancien = _ctx_str(context, "ancien_lieu") or get_lieu_travail(employee_data, {})
        nouveau = _ctx_str(context, "nouveau_lieu")
        _render_table(doc, "Lieu de travail", ancien or "…………………", nouveau or "…………………")
        return

    texte = motif or "Modification contractuelle convenue entre les parties."
    _add_body(doc, "Les parties conviennent de la modification suivante :")
    p = doc.add_paragraph()
    p.add_run(texte).italic = True


def _render_table(doc: Document, label: str, ancien: str, nouveau: str) -> None:
    table = doc.add_table(rows=2, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Élément"
    hdr[1].text = "Ancienne rédaction"
    hdr[2].text = "Nouvelle rédaction"
    row = table.rows[1].cells
    row[0].text = label
    row[1].text = ancien
    row[2].text = nouveau


def generate_avenant_docx(
    employee_data: Dict[str, Any],
    company_data: Dict[str, Any],
    context: Optional[Dict[str, Any]] = None,
) -> bytes:
    """Génère un .docx d'avenant avec le même contenu que le PDF EYWAI."""
    ctx = context or {}
    avenant_type = _ctx_str(ctx, "type_avenant") or "avenant_general"
    objet_label = AVENANT_TYPE_LABELS.get(avenant_type, AVENANT_TYPE_LABELS["avenant_general"])

    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(11.5)
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)

    company_name = get_company_name(company_data)
    company_address = get_company_address(company_data) or "…………………"
    company_siret = safe_str(company_data.get("siret")) or "…………………"
    company_city = get_company_city(company_data) or "…………………"
    signatory, signatory_title = get_company_signatory(company_data)
    if signatory == "Le service RH":
        signatory = "Le représentant légal"
    if not signatory_title:
        signatory_title = "Employeur"

    first_name = safe_str(employee_data.get("first_name", ""))
    last_name = safe_str(employee_data.get("last_name", ""))
    job_title = safe_str(employee_data.get("job_title", "")) or "…………………"
    employee_address = safe_str(get_employee_address(employee_data)) or "…………………"
    contract_type = safe_str(employee_data.get("contract_type", "CDI")) or "CDI"
    contract_date = _fmt_date(
        employee_data.get("date_conclusion_contrat")
        or employee_data.get("date_debut_execution")
        or employee_data.get("hire_date")
    )
    date_effet = _fmt_date(ctx.get("date_effet"))
    date_avenant = _fmt_date(ctx.get("date_avenant") or date.today().isoformat())
    motif = _ctx_str(ctx, "motif_avenant", "motif")
    date_fin = _fmt_date(ctx.get("date_fin_avenant"))

    logo_bytes = resolve_company_logo(company_data)
    if logo_bytes:
        try:
            doc.add_picture(BytesIO(logo_bytes), width=Cm(2.5))
        except Exception:
            pass

    doc.add_paragraph(company_name).runs[0].bold = True
    doc.add_paragraph(f"SIRET : {company_siret}")
    doc.add_paragraph(f"Siège social : {company_address}")
    doc.add_paragraph()

    _add_title(doc, "Avenant au contrat de travail")

    p = doc.add_paragraph()
    p.add_run("Entre les soussignés :").bold = True
    doc.add_paragraph(
        f"{company_name}\nSIRET : {company_siret}\nSiège social : {company_address}\n"
        "Ci-après dénommée « l'Employeur »,"
    )
    c = doc.add_paragraph("D'une part,")
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        f"{first_name} {last_name}\nDomicilié(e) : {employee_address}\n"
        f"Emploi actuel : {job_title}\nCi-après dénommé(e) « le Salarié »,"
    )
    c2 = doc.add_paragraph("D'autre part,")
    c2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_article_title(doc, "Préambule — Rappel du contrat initial")
    _add_body(
        doc,
        f"Le Salarié est lié à l'Employeur par un contrat de travail de type "
        f"{contract_type}, conclu en date du {contract_date}.",
    )
    _add_body(
        doc,
        "Les parties souhaitent modifier certaines dispositions de ce contrat et "
        "conviennent de conclure le présent avenant.",
    )

    _add_article_title(doc, "Article 1 — Objet de l'avenant")
    _add_body(
        doc,
        f"Le présent avenant a pour objet de modifier les dispositions relatives à "
        f"{objet_label} du contrat de travail initial.",
    )
    if motif and avenant_type != "avenant_general":
        _add_body(doc, f"Motif : {motif}")

    _add_article_title(doc, "Article 2 — Modifications apportées")
    _add_modif_table(doc, avenant_type, ctx, employee_data)

    _add_article_title(doc, "Article 3 — Date d'effet")
    _add_body(
        doc,
        f"Les modifications prévues au présent avenant prennent effet à compter du "
        f"{date_effet}.",
    )
    if date_fin and date_fin != "…………………":
        _add_body(
            doc,
            f"Le présent avenant prend fin le {date_fin}, sauf prorogation ou "
            "renouvellement convenu par écrit entre les parties.",
        )

    _add_article_title(doc, "Article 4 — Maintien des autres clauses")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(
        "Toutes les autres clauses du contrat de travail initial non modifiées par le "
        "présent avenant demeurent inchangées et continuent de produire leurs effets."
    )
    run.italic = True

    if motif and _is_economic_motif(motif):
        note = doc.add_paragraph()
        note_run = note.add_run(
            "Note — Modification pour motif économique (art. L1222-6 C. trav.) : la "
            "proposition de modification a été/devra être adressée au salarié par "
            "lettre recommandée avec accusé de réception. Le salarié dispose d'un "
            "délai de réflexion d'un mois à compter de la réception de la "
            "proposition avant toute signature."
        )
        note_run.bold = True
        note_run.font.size = Pt(9.5)

    doc.add_paragraph()
    doc.add_paragraph(
        f"Fait à {company_city}, le {date_avenant}, en deux exemplaires originaux, "
        "remis à chaque partie."
    )

    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=2)
    left, right = table.rows[0].cells
    left.paragraphs[0].add_run(
        f"{signatory}\n{signatory_title}\nSignature de l'Employeur\n"
        "(précédée de la mention manuscrite « Lu et approuvé »)"
    )
    right.paragraphs[0].add_run(
        f"{first_name} {last_name}\nSignature du Salarié\n"
        "(précédée de la mention manuscrite « Lu et approuvé »)"
    )

    out = BytesIO()
    doc.save(out)
    return out.getvalue()
