"""
Génération Word (.docx) de contrat de travail — équivalent éditable du PDF.

Même contenu que contract.py (mentions obligatoires art. L1221-1 et suivants),
mais produit un fichier Word que le service RH peut rouvrir et retravailler
(ex. reprendre un contrat existant comme base pour un nouveau salarié).
"""

from __future__ import annotations

from io import BytesIO
from typing import Any, Dict

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

from app.shared.infrastructure.pdf.helpers import (
    format_periode_essai,
    format_salary_euros,
    get_company_address,
    get_company_city,
    get_company_name,
    get_company_signatory,
    get_convention_collective,
    get_employee_address,
    get_lieu_travail,
    get_salary_amount,
    resolve_company_logo,
    safe_str,
)

from datetime import date


def _fmt_date(value: Any) -> str:
    if not value:
        return "…………………"
    if isinstance(value, str):
        try:
            return date.fromisoformat(value[:10]).strftime("%d/%m/%Y")
        except ValueError:
            return value
    if isinstance(value, date):
        return value.strftime("%d/%m/%Y")
    return str(value)


def _is_cdd(contract_type: str) -> bool:
    ct = contract_type.upper()
    return "CDD" in ct or "DETERMINE" in ct or "DÉTERMINÉ" in ct


def _add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(14)
    p.paragraph_format.space_after = Pt(14)


def _add_article_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11.5)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def _add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text)
    p.paragraph_format.space_after = Pt(4)


def generate_contract_docx(
    employee_data: Dict[str, Any],
    company_data: Dict[str, Any],
) -> bytes:
    """Génère un .docx de contrat de travail avec le même contenu que le PDF EYWAI."""
    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(11.5)
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)

    company_name = get_company_name(company_data)
    company_address = get_company_address(company_data) or "…………………"
    company_siret = safe_str(company_data.get("siret")) or "…………………"
    company_city = get_company_city(company_data) or "…………………"
    signatory, signatory_title = get_company_signatory(company_data)
    if signatory == "Le service RH":
        signatory = "Le représentant légal"
    if not signatory_title:
        signatory_title = "Employeur"

    first_name = safe_str(employee_data.get("first_name", ""))
    last_name = safe_str(employee_data.get("last_name", ""))
    hire_date = _fmt_date(employee_data.get("hire_date"))
    contract_type = safe_str(employee_data.get("contract_type", "CDI")) or "CDI"
    job_title = safe_str(employee_data.get("job_title", "")) or "…………………"
    duree_hebdo = employee_data.get("duree_hebdomadaire", 35)
    is_partiel = employee_data.get("is_temps_partiel") or (
        isinstance(duree_hebdo, (int, float)) and float(duree_hebdo) < 35
    )
    statut = safe_str(employee_data.get("statut", "Non-cadre"))
    lieu_travail = safe_str(get_lieu_travail(employee_data, company_data))
    salaire_str = format_salary_euros(employee_data)
    cc = employee_data.get("classification_conventionnelle") or {}
    groupe = safe_str(cc.get("groupe_emploi")) if isinstance(cc, dict) else ""
    classe = safe_str(cc.get("classe_emploi")) if isinstance(cc, dict) else ""
    coefficient = safe_str(cc.get("coefficient")) if isinstance(cc, dict) else ""
    convention = safe_str(get_convention_collective(company_data, employee_data))
    periode_essai = safe_str(format_periode_essai(employee_data))
    employee_address = safe_str(get_employee_address(employee_data)) or "…………………"
    date_naissance = safe_str(employee_data.get("date_naissance")) or "…………………"
    lieu_naissance = safe_str(employee_data.get("lieu_naissance")) or "…………………"
    nationalite = safe_str(employee_data.get("nationalite", "Française"))
    nir = safe_str(employee_data.get("nir")) or "…………………"

    logo_bytes = resolve_company_logo(company_data)
    if logo_bytes:
        try:
            doc.add_picture(BytesIO(logo_bytes), width=Cm(2.5))
        except Exception:
            pass

    hdr = doc.add_paragraph()
    hdr.add_run(company_name).bold = True
    doc.add_paragraph(f"SIRET : {company_siret}")
    doc.add_paragraph(f"Siège social : {company_address}")
    doc.add_paragraph()

    _add_title(doc, f"Contrat de travail — {contract_type}")

    p = doc.add_paragraph()
    p.add_run("Entre les soussignés :").bold = True
    doc.add_paragraph(
        f"{company_name}\nSIRET : {company_siret}\nSiège social : {company_address}\n"
        "Ci-après dénommée « l'Employeur »,"
    )
    c = doc.add_paragraph("D'une part,")
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        f"{first_name} {last_name}\nNé(e) le {date_naissance} à {lieu_naissance}\n"
        f"Nationalité : {nationalite}\nDomicilié(e) : {employee_address}\n"
        f"N° de sécurité sociale : {nir}\nCi-après dénommé(e) « le Salarié »,"
    )
    c2 = doc.add_paragraph("D'autre part,")
    c2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_title(doc, "Il a été convenu ce qui suit :")

    _add_article_title(doc, "Article 1 — Engagement et date d'entrée")
    _add_body(
        doc,
        f"L'Employeur engage le Salarié, qui accepte, à compter du {hire_date}, "
        f"en qualité de {job_title}, sous contrat de type {contract_type}.",
    )

    _add_article_title(doc, "Article 2 — Fonctions et lieu de travail")
    _add_body(
        doc,
        f"Le Salarié exercera les fonctions de {job_title}, rattaché(e) à la direction "
        "dont il/elle relève.",
    )
    _add_body(
        doc,
        f"Le lieu habituel de travail est fixé à {lieu_travail}. L'Employeur se réserve "
        "la possibilité de modifier le lieu de travail dans le respect des dispositions "
        "légales et conventionnelles.",
    )
    _add_body(
        doc,
        "Le Salarié s'engage à consacrer l'intégralité de son activité professionnelle "
        "à l'Employeur et à exécuter les missions confiées avec diligence et loyauté.",
    )

    _add_article_title(doc, "Article 3 — Durée du contrat")
    if _is_cdd(contract_type):
        fin = _fmt_date(
            employee_data.get("date_fin_contrat")
            or employee_data.get("contract_end_date")
            or employee_data.get("end_date")
        )
        motif = safe_str(
            employee_data.get("motif_cdd") or employee_data.get("cdd_motif")
        ) or "remplacement ou accroissement temporaire d'activité"
        _add_body(
            doc,
            f"Le présent contrat est conclu pour une durée déterminée, du {hire_date} "
            f"au {fin}, pour le motif suivant : {motif}.",
        )
        _add_body(
            doc,
            "Il prendra fin à l'échéance du terme, sous réserve des cas de prorogation "
            "ou de renouvellement prévus par la loi et la convention collective.",
        )
    else:
        _add_body(
            doc,
            "Le présent contrat est conclu pour une durée indéterminée (CDI). "
            "Il prend effet à la date d'entrée en fonction mentionnée à l'article 1.",
        )

    _add_article_title(doc, "Article 4 — Période d'essai")
    _add_body(
        doc,
        f"Le contrat est conclu sous réserve d'une période d'essai de {periode_essai}.",
    )
    _add_body(
        doc,
        "Chacune des parties pourra rompre le contrat pendant cette période, dans les "
        "conditions et délais prévus par le Code du travail et la convention collective "
        "applicable.",
    )

    _add_article_title(doc, "Article 5 — Durée du travail")
    _add_body(
        doc,
        f"Le Salarié est employé à temps {'partiel' if is_partiel else 'complet'}.",
    )
    _add_body(
        doc,
        f"La durée hebdomadaire de travail est fixée à {duree_hebdo} heures, répartie "
        "conformément au planning et aux usages de l'entreprise, dans le respect des "
        "durées maximales légales.",
    )

    _add_article_title(doc, "Article 6 — Rémunération")
    _add_body(
        doc,
        f"En contrepartie de son travail, le Salarié percevra une rémunération "
        f"mensuelle brute de {salaire_str}, pour un temps de travail à hauteur de la "
        "durée prévue à l'article 5.",
    )
    if groupe or classe or coefficient:
        _add_body(
            doc,
            f"Classification conventionnelle : groupe {groupe or '—'}, "
            f"classe {classe or '—'}, coefficient {coefficient or '—'}.",
        )
    _add_body(
        doc,
        "Le salaire est payable mensuellement, par virement bancaire, au plus tard le "
        "dernier jour ouvré du mois.",
    )

    _add_article_title(doc, "Article 7 — Congés payés et absences")
    _add_body(
        doc,
        "Le Salarié bénéficie des congés payés dans les conditions prévues par le Code "
        "du travail et la convention collective. Les absences doivent être signalées "
        "et justifiées selon les règles internes de l'entreprise.",
    )

    _add_article_title(doc, "Article 8 — Convention collective et statut")
    _add_body(doc, f"Le Salarié relève du statut {statut}.")
    _add_body(
        doc,
        "Le présent contrat est soumis aux dispositions du Code du travail, à la "
        f"convention collective {convention}, ainsi qu'au règlement intérieur et aux "
        "usages en vigueur dans l'entreprise.",
    )

    _add_article_title(doc, "Article 9 — Protection sociale complémentaire")
    _add_body(
        doc,
        "Le Salarié bénéficie, dans les conditions prévues par la loi et les accords "
        "applicables, de la couverture complémentaire santé (mutuelle) et de la "
        "prévoyance mises en place par l'Employeur. Les cotisations afférentes sont "
        "réparties conformément aux textes en vigueur.",
    )

    _add_article_title(doc, "Article 10 — Obligations professionnelles")
    _add_body(
        doc,
        "Le Salarié est tenu au respect du règlement intérieur, des consignes de "
        "sécurité, du secret professionnel et de la confidentialité des informations "
        "dont il/elle a connaissance dans l'exercice de ses fonctions.",
    )
    _add_body(
        doc,
        "Le Salarié s'engage à respecter les règles relatives à la protection des "
        "données à caractère personnel (RGPD) applicables dans l'entreprise.",
    )

    _add_article_title(doc, "Article 11 — Rupture du contrat")
    _add_body(
        doc,
        "Le contrat pourra être rompu par l'une ou l'autre des parties dans les "
        "conditions et formes prévues par le Code du travail, notamment en cas de "
        "démission, licenciement ou rupture conventionnelle, sous réserve du respect "
        "des délais de préavis légaux ou conventionnels.",
    )

    mention = doc.add_paragraph()
    mention.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    mention_run = mention.add_run(
        "Il contient les mentions obligatoires prévues à l'article L1221-1 du Code du "
        "travail."
        + (
            " Le salaire indiqué est provisoire et devra être complété avant signature "
            "définitive."
            if get_salary_amount(employee_data) <= 0
            else ""
        )
    )
    mention_run.font.size = Pt(9.5)

    doc.add_paragraph()
    doc.add_paragraph(
        f"Fait à {company_city}, le {hire_date}, en deux exemplaires originaux, "
        "remis à chaque partie."
    )

    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=2)
    left, right = table.rows[0].cells
    left.paragraphs[0].add_run(
        f"{signatory}\n{signatory_title}\nSignature de l'Employeur\n"
        "(précédée de la mention « Lu et approuvé »)"
    )
    right.paragraphs[0].add_run(
        f"{first_name} {last_name}\nSignature du Salarié\n"
        "(précédée de la mention « Lu et approuvé »)"
    )

    out = BytesIO()
    doc.save(out)
    return out.getvalue()
