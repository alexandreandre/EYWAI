#!/usr/bin/env python3
"""
Prépare (variables EYWAI) et uploade les trames contrat Comitech Composite.

Sources : fichiers Word sur le Bureau de l'utilisateur (ou --source-dir).
Cibles : bucket Supabase document_templates + bibliothèque entreprise.

Usage (depuis backend/, venv activé) :
  python scripts/upload_comitech_contract_templates.py --prepare-only
  python scripts/upload_comitech_contract_templates.py --upload
  python scripts/upload_comitech_contract_templates.py --prepare-only --upload
  python scripts/upload_comitech_contract_templates.py --dry-run --upload
"""

from __future__ import annotations

import argparse
import io
import os
import re
import sys
from pathlib import Path
from typing import Any, Callable, Iterable

BACKEND_ROOT = Path(__file__).resolve().parents[1]
if str(BACKEND_ROOT) not in sys.path:
    sys.path.insert(0, str(BACKEND_ROOT))

env_file = BACKEND_ROOT / ".env"
if env_file.exists():
    for line in env_file.read_text().splitlines():
        line = line.strip()
        if line and not line.startswith("#") and "=" in line:
            k, v = line.split("=", 1)
            os.environ.setdefault(k.strip(), v.strip().strip('"'))

COMPANY_SIRET = "49861035100013"
COMPANY_NAME = "Comitech Composite"
DEFAULT_SOURCE_DIR = Path.home() / "Desktop"
OUTPUT_DIR = (
    BACKEND_ROOT / "app" / "static" / "document_templates" / "comitech_composite"
)

ELLIPSIS = "…"
DOTS = re.compile(rf"(?:{re.escape(ELLIPSIS)}|\.)+")
FILL = r"[…\.]+"

Replacement = tuple[str | re.Pattern[str], str]

TEMPLATE_SPECS: list[dict[str, Any]] = [
    {
        "source_name": "CDD temps plein NC.docx",
        "output_name": "CDD_temps_plein_NC.docx",
        "document_type": "cdd",
        "name": "CDD temps plein NC",
        "is_default": True,
    },
    {
        "source_name": "CDI temps plein NC.docx",
        "output_name": "CDI_temps_plein_NC.docx",
        "document_type": "cdi",
        "name": "CDI temps plein NC",
        "is_default": True,
    },
    {
        "source_name": "CDI temps plein NC - Forfait jour.docx",
        "output_name": "CDI_temps_plein_NC_forfait_jour.docx",
        "document_type": "cdi",
        "name": "CDI NC forfait jour",
        "is_default": False,
    },
    {
        "source_name": "CDI temps plein cadre  - Forfait jour.docx",
        "output_name": "CDI_temps_plein_cadre_forfait_jour.docx",
        "document_type": "cdi",
        "name": "CDI cadre forfait jour",
        "is_default": False,
    },
    {
        "source_name": "Avenant renouvellement 1 de CDD.docx",
        "output_name": "Avenant_renouvellement_CDD_1.docx",
        "document_type": "avenant_general",
        "name": "Renouvellement CDD — 1er",
        "is_default": False,
    },
    {
        "source_name": "Avenant renouvellement 2 de CDD.docx",
        "output_name": "Avenant_renouvellement_CDD_2.docx",
        "document_type": "avenant_general",
        "name": "Renouvellement CDD — 2e",
        "is_default": False,
    },
    {
        "source_name": "Avenant passage en CDI NC.docx",
        "output_name": "Avenant_passage_CDI_NC.docx",
        "document_type": "avenant_general",
        "name": "Passage CDD → CDI NC",
        "is_default": False,
    },
]


def _common_replacements() -> list[Replacement]:
    return [
        ("COMITECH", "{{nom_entreprise}}"),
        ("498 610 351 00013", "{{siret}}"),
        ("49861035100013", "{{siret}}"),
        ("code NAF 2229 A", "code NAF {{code_ape}}"),
        ("Z.A.la Pelissière - 01300 BELLEY", "{{adresse_entreprise}}"),
        (
            "827000002161193744",
            "{{urssaf_number}}",
        ),
        ("Monsieur Gérault VERNY", "Monsieur {{nom_signataire_rh}}"),
        (
            "en sa qualité de Président.La société",
            "en sa qualité de {{qualite_signataire_rh}}.",
        ),
        (
            "en sa qualité de Président.",
            "en sa qualité de {{qualite_signataire_rh}}.",
        ),
        ("Monsieur MARTIN Martin", "Monsieur {{nom}} {{prenom}}"),
        ("Monsieur Martin Martin", "Monsieur {{nom}} {{prenom}}"),
        (
            "Né le ………………… à …………………. (……………….) de nationalité …………………",
            "Né(e) le {{date_naissance}} à {{lieu_naissance}} de nationalité {{nationalite}}",
        ),
        (
            re.compile(
                rf"Né le {FILL}+ à {FILL}+ ?\({FILL}+\) de nationalité {FILL}+"
            ),
            "Né(e) le {{date_naissance}} à {{lieu_naissance}} de nationalité {{nationalite}}",
        ),
        (
            re.compile(
                rf"Dont le n° de Sécurité Sociale est\s*:?\s*{FILL}+"
            ),
            "Dont le n° de Sécurité Sociale est : {{numero_securite_sociale}}",
        ),
        (
            "Porteur du titre de séjour N° ……………………… valable du ……………….. au …………………",
            "Porteur du titre de séjour N° {{numero_titre_sejour}} valable jusqu'au {{titre_sejour_fin}}",
        ),
        (
            re.compile(
                rf"Porteur du titre de séjour N° {FILL}+ valable du {FILL}+ au {FILL}+"
            ),
            "Porteur du titre de séjour N° {{numero_titre_sejour}} valable jusqu'au {{titre_sejour_fin}}",
        ),
        ("Fait à Belley,", "Fait à {{signature_lieu}},"),
        ("Le X", "Le {{signature_date}}"),
        (
            "Convention collective nationale de la Plasturgie",
            "{{convention_collective}}",
        ),
    ]


def _contract_replacements() -> list[Replacement]:
    return _common_replacements() + [
        (
            re.compile(
                rf"engagé à compter du {FILL}+ en contrat de travail à durée indéterminée, en qualité de {FILL}+, coefficient {FILL}+"
            ),
            "engagé à compter du {{date_debut_contrat}} en contrat de travail à durée indéterminée, en qualité de {{poste}}, coefficient {{coefficient}}",
        ),
        (
            re.compile(
                rf"engagé à compter du {FILL}+ en contrat de travail à durée indéterminée, en qualité d'{FILL}+, coefficient {FILL}+"
            ),
            "engagé à compter du {{date_debut_contrat}} en contrat de travail à durée indéterminée, en qualité de {{poste}}, coefficient {{coefficient}}",
        ),
        (
            "engagé à compter du ………………… en contrat de travail à durée indéterminée, en qualité d'Agent de production, coefficient ………",
            "engagé à compter du {{date_debut_contrat}} en contrat de travail à durée indéterminée, en qualité de {{poste}}, coefficient {{coefficient}}",
        ),
        (
            "engagé à compter du ………………… en contrat de travail à durée indéterminée, en qualité d’Agent de production, coefficient ………, conformément à la convention collective applicable.",
            "engagé à compter du {{date_debut_contrat}} en contrat de travail à durée indéterminée, en qualité de {{poste}}, coefficient {{coefficient}}, conformément à la convention collective applicable.",
        ),
        (
            re.compile(
                rf"engagé à compter du {FILL}+ en contrat de travail à durée indéterminée, en qualité d.{{1}}Agent de production, coefficient {FILL}+"
            ),
            "engagé à compter du {{date_debut_contrat}} en contrat de travail à durée indéterminée, en qualité de {{poste}}, coefficient {{coefficient}}",
        ),
        (
            re.compile(
                rf"engagé à compter du {FILL}+ en contrat de travail à durée indéterminée, en qualité de {FILL}+, coefficient {FILL}+"
            ),
            "engagé à compter du {{date_debut_contrat}} en contrat de travail à durée indéterminée, en qualité de {{poste}}, coefficient {{coefficient}}",
        ),
        (
            re.compile(rf"rémunération mensuelle brute de {FILL}+ euros"),
            "rémunération mensuelle brute de {{salaire_brut_mensuel}}",
        ),
        (
            "Né le ………………… à …………………. (………….) de nationalité …………………",
            "Né(e) le {{date_naissance}} à {{lieu_naissance}} de nationalité {{nationalite}}",
        ),
        (
            "Porteur du titre de séjour N° ……………………… valable du ………………. au …………………",
            "Porteur du titre de séjour N° {{numero_titre_sejour}} valable jusqu'au {{titre_sejour_fin}}",
        ),
        (
            re.compile(rf"période d.essai de {FILL}+ mois"),
            "période d'essai de {{periode_essai_duree}}",
        ),
        (
            "Le présent contrat est assorti d’une période d’essai de ………………… mois",
            "Le présent contrat est assorti d'une période d'essai de {{periode_essai_duree}}",
        ),
        ("à compter du 1er avril 2026,", "à compter du {{date_debut_contrat}},"),
        (
            "à compter du 1er avril 2026 jusqu’au 30 avril 2026",
            "à compter du {{date_debut_contrat}} jusqu'au {{date_fin_contrat}}",
        ),
        (
            "Le contrat à durée déterminée est conclu à compter du 1er avril 2026 jusqu’au 30 avril 2026.",
            "Le contrat à durée déterminée est conclu à compter du {{date_debut_contrat}} jusqu'au {{date_fin_contrat}}.",
        ),
        (
            "en qualité d’Agent de production, coefficient 700",
            "en qualité de {{poste}}, coefficient {{coefficient}}",
        ),
        (
            re.compile(
                r"{{nom}} {{prenom}} est engagé dans le cadre d.un contrat à durée déterminée à compter du 1er avril 2026,"
            ),
            "Le Salarié est engagé dans le cadre d'un contrat à durée déterminée à compter du {{date_debut_contrat}},",
        ),
        (
            "rémunération mensuelle brute de 2000 euros",
            "rémunération mensuelle brute de {{salaire_brut_mensuel}}",
        ),
        (
            "fixée à 2 055,51 euros bruts mensuels",
            "fixée à {{salaire_brut_mensuel}}",
        ),
        (
            "Monsieur Martin Martin est engagé dans le cadre d'un contrat à durée déterminée à compter du 1er avril 2026,\nen qualité d'Agent de production, coefficient 700",
            "Le Salarié est engagé dans le cadre d'un contrat à durée déterminée à compter du {{date_debut_contrat}}, en qualité de {{poste}}, coefficient {{coefficient}}",
        ),
        (
            "Monsieur Martin Martin est engagé dans le cadre d'un contrat à durée déterminée à compter du 1er avril 2026, en qualité d'Agent de production, coefficient 700",
            "Le Salarié est engagé dans le cadre d'un contrat à durée déterminée à compter du {{date_debut_contrat}}, en qualité de {{poste}}, coefficient {{coefficient}}",
        ),
        (
            re.compile(
                r"Monsieur Martin Martin est engagé dans le cadre d'un contrat à durée déterminée à compter du .+?, en qualité d'Agent de production, coefficient 700"
            ),
            "Le Salarié est engagé dans le cadre d'un contrat à durée déterminée à compter du {{date_debut_contrat}}, en qualité de {{poste}}, coefficient {{coefficient}}",
        ),
        (
            "La période d'essai prendra fin le 4 avril 2026 au soir.",
            "La période d'essai prendra fin le {{fin_periode_essai}} au soir.",
        ),
        (
            "term fixé au 30 avril 2026",
            "term fixé au {{date_fin_contrat}}",
        ),
        (
            "terminera de plein droit à l'échéance du terme fixé au 30 avril 2026",
            "terminera de plein droit à l'échéance du terme fixé au {{date_fin_contrat}}",
        ),
        (
            "prendra fin de plein droit à l'échéance du terme fixé au 30 avril 2026",
            "prendra fin de plein droit à l'échéance du terme fixé au {{date_fin_contrat}}",
        ),
        (
            "Le Salarié exercera ses fonctions à Belley",
            "Le Salarié exercera ses fonctions à {{lieu_travail}}",
        ),
    ]


def _avenant_replacements() -> list[Replacement]:
    return _common_replacements() + [
        (
            "Le contrat de travail à durée déterminée conclu entre les parties le ………………………..est transformé",
            "Le contrat de travail à durée déterminée conclu entre les parties le {{date_debut_contrat}} est transformé",
        ),
        (
            re.compile(
                rf"contrat de travail à durée déterminée conclu entre les parties le {FILL}+est transformé"
            ),
            "contrat de travail à durée déterminée conclu entre les parties le {{date_debut_contrat}} est transformé",
        ),
        (
            "En conséquence, la date d'ancienneté du Salarié est fixée au ……………………………………….",
            "En conséquence, la date d'ancienneté du Salarié est fixée au {{date_debut_contrat}}.",
        ),
        (
            re.compile(
                rf"date d.ancienneté du Salarié est fixée au {FILL}+"
            ),
            "date d'ancienneté du Salarié est fixée au {{date_debut_contrat}}",
        ),
        (
            "À compter du …………………….., la rémunération du Salarié est fixée à 2 055,51 euros bruts mensuels",
            "À compter du {{date_effet}}, la rémunération du Salarié est fixée à {{salaire_brut_mensuel}}",
        ),
        (
            re.compile(
                rf"À compter du {FILL}+, la rémunération du Salarié est fixée à .+? euros bruts mensuels"
            ),
            "À compter du {{date_effet}}, la rémunération du Salarié est fixée à {{salaire_brut_mensuel}}",
        ),
        (
            re.compile(
                rf"en contrat de travail à durée indéterminée à compter du {FILL}+"
            ),
            "en contrat de travail à durée indéterminée à compter du {{date_effet}}",
        ),
        (
            re.compile(
                rf"date d'ancienneté du Salarié est fixée au {FILL}+"
            ),
            "date d'ancienneté du Salarié est fixée au {{date_debut_contrat}}",
        ),
        (
            re.compile(rf"À compter du {FILL}+, la rémunération du Salarié est fixée à .+? mensuels"),
            "À compter du {{date_effet}}, la rémunération du Salarié est fixée à {{salaire_brut_mensuel}}",
        ),
        (
            re.compile(
                rf"contrat de travail à durée déterminée conclu entre les parties en date du {FILL}+ est renouvelé"
            ),
            "contrat de travail à durée déterminée conclu entre les parties en date du {{date_debut_contrat}} est renouvelé",
        ),
        (
            re.compile(rf"terme du contrat est fixé au {FILL}+ et celui-ci est prolongé pour la période courant du {FILL}+ au {FILL}+"),
            "terme du contrat est fixé au {{date_fin_contrat}} et celui-ci est prolongé pour la période courant du {{date_effet}} au {{date_fin_contrat}}",
        ),
    ]


def _replacements_for_spec(spec: dict[str, Any]) -> list[Replacement]:
    dtype = spec["document_type"]
    if dtype == "avenant_general":
        return _avenant_replacements()
    return _contract_replacements()


def _apply_replacements(text: str, replacements: list[Replacement]) -> str:
    out = text
    for old, new in replacements:
        if isinstance(old, re.Pattern):
            out = old.sub(new, out)
        else:
            out = out.replace(old, new)
    return out


def _replace_in_paragraph(paragraph, replacements: list[Replacement]) -> bool:
    original = paragraph.text
    if not original:
        return False
    updated = _apply_replacements(original, replacements)
    if updated == original:
        return False
    if paragraph.runs:
        paragraph.runs[0].text = updated
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(updated)
    return True


def _iter_paragraphs(doc) -> Iterable[Any]:
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
    for section in doc.sections:
        for header_footer in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
        ):
            if header_footer is None:
                continue
            for paragraph in header_footer.paragraphs:
                yield paragraph
            for table in header_footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            yield paragraph


def prepare_docx(source: Path, destination: Path, replacements: list[Replacement]) -> dict[str, int]:
    from docx import Document

    doc = Document(str(source))
    changed = 0
    for paragraph in _iter_paragraphs(doc):
        if _replace_in_paragraph(paragraph, replacements):
            changed += 1
    destination.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(destination))
    return {"paragraphs_changed": changed}


def _find_company_id(supabase) -> str:
    r = (
        supabase.table("companies")
        .select("id, company_name")
        .eq("siret", COMPANY_SIRET)
        .maybe_single()
        .execute()
    )
    if r and r.data:
        return str(r.data["id"])
    r2 = (
        supabase.table("companies")
        .select("id, company_name")
        .ilike("company_name", COMPANY_NAME)
        .limit(1)
        .execute()
    )
    rows = r2.data or []
    if rows:
        return str(rows[0]["id"])
    raise RuntimeError(
        f"Entreprise {COMPANY_NAME} introuvable — lancer setup_comitech_composite.py d'abord."
    )


def _find_template_by_name(
    supabase, company_id: str, document_type: str, name: str
) -> dict | None:
    r = (
        supabase.table("document_templates")
        .select("*")
        .eq("company_id", company_id)
        .eq("document_type", document_type)
        .eq("name", name)
        .eq("status", "active")
        .limit(1)
        .execute()
    )
    rows = r.data or []
    return rows[0] if rows else None


def upload_template(
    supabase,
    company_id: str,
    spec: dict[str, Any],
    file_path: Path,
    *,
    dry_run: bool,
) -> dict[str, Any]:
    from app.modules.document_library.application.commands import (
        create_template,
        upload_template_file,
        validate_template_bytes,
    )
    from app.modules.document_library.infrastructure.repository import (
        document_library_repository,
    )
    from app.modules.document_library.schemas.requests import (
        DocumentTemplateCreate,
        DocumentTemplateUpdate,
    )

    file_bytes = file_path.read_bytes()
    validation = validate_template_bytes(file_bytes, file_path.name)
    unknown = validation.get("unknown_variables") or []

    existing = _find_template_by_name(
        supabase, company_id, spec["document_type"], spec["name"]
    )

    if dry_run:
        action = "dry_run_update" if existing else "dry_run_create"
        return {
            "name": spec["name"],
            "document_type": spec["document_type"],
            "action": action,
            "unknown_variables": unknown,
            "file": str(file_path),
        }

    if existing:
        template_id = str(existing["id"])
        action = "updated"
    else:
        created = create_template(
            company_id,
            DocumentTemplateCreate(
                document_type=spec["document_type"],
                name=spec["name"],
            ),
            created_by=None,
        )
        template_id = str(created["id"])
        action = "created"

    version = upload_template_file(
        company_id,
        template_id,
        file_bytes,
        file_path.name,
        created_by=None,
    )

    if spec.get("is_default"):
        document_library_repository.update(
            template_id,
            company_id,
            DocumentTemplateUpdate(is_default=True),
        )

    return {
        "name": spec["name"],
        "document_type": spec["document_type"],
        "template_id": template_id,
        "version": version.get("version"),
        "action": action,
        "unknown_variables": unknown,
        "is_default": bool(spec.get("is_default")),
    }


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Prépare et uploade les trames contrat Comitech Composite"
    )
    parser.add_argument(
        "--source-dir",
        type=Path,
        default=DEFAULT_SOURCE_DIR,
        help="Dossier contenant les 7 fichiers Word sources",
    )
    parser.add_argument(
        "--prepare-only",
        action="store_true",
        help="Génère les .docx préparés dans app/static/document_templates/comitech_composite/",
    )
    parser.add_argument(
        "--upload",
        action="store_true",
        help="Upload vers la bibliothèque Supabase Comitech Composite",
    )
    parser.add_argument("--dry-run", action="store_true")
    args = parser.parse_args()

    if not args.prepare_only and not args.upload:
        args.prepare_only = True
        args.upload = True

    prepared: list[dict[str, Any]] = []

    if args.prepare_only:
        print(f"Préparation des trames depuis {args.source_dir} → {OUTPUT_DIR}")
        for spec in TEMPLATE_SPECS:
            source = args.source_dir / spec["source_name"]
            if not source.is_file():
                print(f"  ERREUR : fichier source absent : {source}")
                return 1
            dest = OUTPUT_DIR / spec["output_name"]
            stats = prepare_docx(source, dest, _replacements_for_spec(spec))
            prepared.append({"spec": spec, "path": dest, **stats})
            print(
                f"  ✓ {spec['name']} — {stats['paragraphs_changed']} paragraphe(s) modifié(s) → {dest.name}"
            )

    upload_paths: list[tuple[dict[str, Any], Path]] = []
    if args.upload:
        for spec in TEMPLATE_SPECS:
            path = OUTPUT_DIR / spec["output_name"]
            if not path.is_file():
                if args.prepare_only:
                    continue
                print(f"  ERREUR : fichier préparé absent : {path} (lancer --prepare-only)")
                return 1
            upload_paths.append((spec, path))

    if args.upload and upload_paths:
        from app.core.database import supabase

        company_id = _find_company_id(supabase)
        print(f"\nUpload bibliothèque documents — {COMPANY_NAME} ({company_id})")
        results = []
        for spec, path in upload_paths:
            result = upload_template(
                supabase, company_id, spec, path, dry_run=args.dry_run
            )
            results.append(result)
            unknown = result.get("unknown_variables") or []
            warn = f" — variables inconnues: {unknown}" if unknown else ""
            print(
                f"  ✓ {result['name']} ({result['document_type']}) "
                f"→ {result['action']}{warn}"
            )
        if args.dry_run:
            print("\n[dry-run] Aucune écriture Supabase effectuée.")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
